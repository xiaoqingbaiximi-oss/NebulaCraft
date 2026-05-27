"""
文件处理 API - 增强版
支持: 多格式解析 / 批量上传 / 文件预览 / 代码高亮 / 文档转换
"""
import io
import os
import time
import base64
from server.config import OUTPUT_DIR, DATA_DIR
from server.utils.security import sanitize_filename

UPLOAD_DIR = os.path.join(DATA_DIR, "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 支持的代码文件扩展名
CODE_EXTENSIONS = {
    "py": "python", "js": "javascript", "ts": "typescript",
    "html": "html", "css": "css", "json": "json",
    "java": "java", "cpp": "cpp", "c": "c", "go": "go",
    "rs": "rust", "rb": "ruby", "php": "php", "sql": "sql",
    "sh": "bash", "yaml": "yaml", "yml": "yaml", "xml": "xml",
    "md": "markdown", "toml": "toml", "ini": "ini",
}


def handle_analyze(file_data, filename):
    """智能文件分析"""
    ext = filename.split(".")[-1].lower() if "." in filename else ""
    
    result = {
        "filename": filename,
        "type": ext or "未知",
        "size": len(file_data),
        "size_formatted": _format_size(len(file_data)),
    }
    
    try:
        # 文本文件
        if ext in ("txt", "md", "csv", "log", "cfg", "conf", "env", "gitignore"):
            text = file_data.decode("utf-8", "ignore")
            lines = text.split("\n")
            result.update({
                "text": text[:10000],
                "lines": len(lines),
                "words": sum(len(l.split()) for l in lines),
                "chars": len(text),
            })
        
        # 代码文件
        elif ext in CODE_EXTENSIONS:
            text = file_data.decode("utf-8", "ignore")
            lines = text.split("\n")
            result.update({
                "text": text[:10000],
                "language": CODE_EXTENSIONS.get(ext, "plaintext"),
                "lines": len(lines),
                "code_lines": sum(1 for l in lines if l.strip() and not l.strip().startswith(("#", "//", "/*", "*", "*/"))),
                "comment_lines": sum(1 for l in lines if l.strip().startswith(("#", "//", "/*", "*", "*/"))),
            })
        
        # JSON
        elif ext == "json":
            text = file_data.decode("utf-8", "ignore")
            try:
                parsed = __import__("json").loads(text)
                result["parsed"] = True
                result["keys"] = list(parsed.keys()) if isinstance(parsed, dict) else None
                result["text"] = __import__("json").dumps(parsed, ensure_ascii=False, indent=2)[:5000]
            except:
                result["text"] = text[:5000]
        
        # 图片文件
        elif ext in ("png", "jpg", "jpeg", "gif", "webp", "bmp", "svg", "ico"):
            from PIL import Image
            img = Image.open(io.BytesIO(file_data))
            result.update({
                "image_size": f"{img.width}x{img.height}",
                "image_mode": img.mode,
                "image_format": img.format,
                "aspect_ratio": f"{img.width / img.height:.2f}",
                "is_animated": getattr(img, "is_animated", False),
                "frames": getattr(img, "n_frames", 1),
            })
            # 生成缩略图
            thumb = img.copy()
            thumb.thumbnail((200, 200), Image.LANCZOS)
            buf = io.BytesIO()
            thumb.save(buf, format="PNG")
            result["thumbnail"] = f"data:image/png;base64,{base64.b64encode(buf.getvalue()).decode()}"
        
        # PDF
        elif ext == "pdf":
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(file_data))
                text = " ".join([p.extract_text() or "" for p in reader.pages[:10]])
                result.update({
                    "text": text[:10000],
                    "pages": len(reader.pages),
                })
            except ImportError:
                result["text"] = "需要 PyPDF2 库解析 PDF"
            except Exception as e:
                result["error"] = f"PDF解析失败: {e}"
        
        # Word
        elif ext in ("docx", "doc"):
            try:
                from docx import Document
                doc = Document(io.BytesIO(file_data))
                text = "\n".join([p.text for p in doc.paragraphs[:50]])
                result.update({
                    "text": text[:10000],
                    "paragraphs": len(doc.paragraphs),
                })
            except ImportError:
                result["text"] = "需要 python-docx 库解析 Word"
            except Exception as e:
                result["error"] = f"Word解析失败: {e}"
        
        # Excel
        elif ext in ("xlsx", "xls", "csv"):
            try:
                import pandas as pd
                if ext == "csv":
                    df = pd.read_csv(io.BytesIO(file_data))
                else:
                    df = pd.read_excel(io.BytesIO(file_data))
                result.update({
                    "text": df.head(100).to_string(),
                    "rows": len(df),
                    "columns": list(df.columns),
                    "shape": str(df.shape),
                })
            except ImportError:
                result["text"] = "需要 pandas 库解析表格"
            except Exception as e:
                result["error"] = f"表格解析失败: {e}"
        
        # 音频文件
        elif ext in ("mp3", "wav", "ogg", "flac", "m4a"):
            result.update({
                "text": f"[音频文件: {ext.upper()}]\n大小: {_format_size(len(file_data))}",
                "media_type": "audio",
            })
        
        # 视频文件
        elif ext in ("mp4", "avi", "mov", "mkv", "webm"):
            result.update({
                "text": f"[视频文件: {ext.upper()}]\n大小: {_format_size(len(file_data))}",
                "media_type": "video",
            })
        
        # 压缩文件
        elif ext in ("zip", "rar", "7z", "tar", "gz"):
            result.update({
                "text": f"[压缩文件: {ext.upper()}]\n大小: {_format_size(len(file_data))}",
                "archive_type": ext,
            })
        
        else:
            result["text"] = f"不支持的文件类型: .{ext}\n大小: {_format_size(len(file_data))}"
    
    except UnicodeDecodeError:
        result["text"] = "[二进制文件，无法预览文本内容]"
    except Exception as e:
        result["error"] = str(e)
    
    return {"ok": True, **result}


def handle_upload(file_data, filename):
    """文件上传保存"""
    safe = sanitize_filename(filename)
    
    # 添加时间戳避免重名
    name, ext = os.path.splitext(safe)
    timestamp = int(time.time() * 1000)
    safe = f"{name}_{timestamp}{ext}"
    
    filepath = os.path.join(UPLOAD_DIR, safe)
    with open(filepath, "wb") as f:
        f.write(file_data)
    
    return {
        "ok": True,
        "filename": safe,
        "path": filepath,
        "size": len(file_data),
        "size_formatted": _format_size(len(file_data)),
    }


def handle_list():
    """列出已上传文件"""
    files = []
    if os.path.exists(UPLOAD_DIR):
        for f in sorted(os.listdir(UPLOAD_DIR), reverse=True):
            fp = os.path.join(UPLOAD_DIR, f)
            if os.path.isfile(fp):
                stat = os.stat(fp)
                files.append({
                    "name": f,
                    "size": stat.st_size,
                    "size_formatted": _format_size(stat.st_size),
                    "time": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(stat.st_mtime)),
                })
    
    return {"ok": True, "files": files[:50]}


def handle_delete(filename):
    """删除上传文件"""
    safe = sanitize_filename(filename)
    fp = os.path.join(UPLOAD_DIR, safe)
    
    if os.path.exists(fp):
        os.remove(fp)
        return {"ok": True, "message": f"已删除: {safe}"}
    return {"ok": False, "error": "文件不存在"}


def _format_size(bytes_val):
    """格式化文件大小"""
    if bytes_val < 1024:
        return f"{bytes_val} B"
    elif bytes_val < 1024 * 1024:
        return f"{bytes_val / 1024:.1f} KB"
    elif bytes_val < 1024 * 1024 * 1024:
        return f"{bytes_val / (1024 * 1024):.1f} MB"
    else:
        return f"{bytes_val / (1024 * 1024 * 1024):.2f} GB"