"""导出 API - PDF / Word / Markdown"""
import os
import io
import time
from server.config import OUTPUT_DIR


def handle(body):
    """导出路由"""
    fmt = body.get("format", "md")
    content = body.get("content", "")
    title = body.get("title", "NebulaCraft 导出")
    
    if not content:
        return {"ok": False, "error": "内容不能为空"}
    
    if fmt == "pdf":
        return _export_pdf(content, title)
    elif fmt == "docx":
        return _export_docx(content, title)
    elif fmt == "html":
        return _export_html(content, title)
    else:
        return _export_md(content, title)


def _export_pdf(content, title):
    """导出 PDF"""
    try:
        # 尝试使用 fpdf2
        try:
            from fpdf import FPDF
        except ImportError:
            return {"ok": False, "error": "请安装 fpdf2: pip install fpdf2"}
        
        pdf = FPDF()
        pdf.add_page()
        
        # 使用内置字体（支持中文有限）
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, title, ln=True, align="C")
        pdf.ln(10)
        
        pdf.set_font("Helvetica", "", 11)
        for line in content.split("\n"):
            # 简单处理 Markdown 标题
            line = line.replace("#", "").replace("*", "").strip()
            if line:
                try:
                    pdf.multi_cell(0, 7, line[:120])
                except:
                    pdf.multi_cell(0, 7, line.encode("ascii", "ignore").decode()[:120])
        
        buf = io.BytesIO()
        pdf.output(buf)
        
        file_id = f"export_{int(time.time() * 1000)}.pdf"
        filepath = os.path.join(OUTPUT_DIR, file_id)
        with open(filepath, "wb") as f:
            f.write(buf.getvalue())
        
        return {"ok": True, "url": f"/output/{file_id}", "format": "pdf"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _export_docx(content, title):
    """导出 Word"""
    try:
        try:
            from docx import Document
        except ImportError:
            return {"ok": False, "error": "请安装 python-docx"}
        
        doc = Document()
        doc.add_heading(title, 0)
        
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            else:
                doc.add_paragraph(line[:500])
        
        buf = io.BytesIO()
        doc.save(buf)
        
        file_id = f"export_{int(time.time() * 1000)}.docx"
        filepath = os.path.join(OUTPUT_DIR, file_id)
        with open(filepath, "wb") as f:
            f.write(buf.getvalue())
        
        return {"ok": True, "url": f"/output/{file_id}", "format": "docx"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _export_html(content, title):
    """导出 HTML"""
    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{ font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; background: #fff; color: #333; line-height: 1.8; }}
        h1 {{ color: #6c5ce7; border-bottom: 2px solid #6c5ce7; padding-bottom: 10px; }}
        h2 {{ color: #a78bfa; margin-top: 24px; }}
        pre {{ background: #f5f5f5; padding: 16px; border-radius: 8px; white-space: pre-wrap; }}
        code {{ background: #f0f0f0; padding: 2px 6px; border-radius: 4px; }}
        .footer {{ text-align: center; color: #999; margin-top: 40px; font-size: .8rem; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <div class="content">
"""
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            html += "<br>\n"
        elif line.startswith("## "):
            html += f"<h2>{line[3:]}</h2>\n"
        elif line.startswith("# "):
            html += f"<h1>{line[2:]}</h1>\n"
        else:
            html += f"<p>{line}</p>\n"
    
    html += f"""</div>
    <div class="footer">导出时间: {time.strftime('%Y-%m-%d %H:%M:%S')} · NebulaCraft v7.0</div>
</body>
</html>"""
    
    file_id = f"export_{int(time.time() * 1000)}.html"
    filepath = os.path.join(OUTPUT_DIR, file_id)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html)
    
    return {"ok": True, "url": f"/output/{file_id}", "format": "html"}


def _export_md(content, title):
    """导出 Markdown"""
    md = f"# {title}\n\n"
    md += f"> 导出时间: {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n"
    md += content
    
    file_id = f"export_{int(time.time() * 1000)}.md"
    filepath = os.path.join(OUTPUT_DIR, file_id)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(md)
    
    return {"ok": True, "url": f"/output/{file_id}", "format": "md"}